import logging
from pathlib import Path
from typing import Dict, List

logger = logging.getLogger(__name__)


def parse_txt(file_bytes: bytes, filename: str) -> List[Dict]:
    """Parse plain text / markdown files."""
    try:
        text = file_bytes.decode("utf-8", errors="replace").strip()
        if not text:
            return []
        # Split into ~page-sized chunks for page numbering
        lines = text.split("\n")
        pages, current, page_num = [], [], 1
        for line in lines:
            current.append(line)
            if len("\n".join(current)) > 3000:
                pages.append({"page": page_num, "text": "\n".join(current).strip()})
                current = []
                page_num += 1
        if current:
            pages.append({"page": page_num, "text": "\n".join(current).strip()})
        return [p for p in pages if p["text"]]
    except Exception as e:
        logger.error(f"TXT parse error for {filename}: {e}")
        return []


def parse_docx(file_bytes: bytes, filename: str) -> List[Dict]:
    """Parse DOCX files using python-docx."""
    import io
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)
        if not full_text:
            return []
        # Group into pages of ~3000 chars
        pages, current, page_num = [], [], 1
        for line in full_text:
            current.append(line)
            if len("\n".join(current)) > 3000:
                pages.append({"page": page_num, "text": "\n".join(current)})
                current = []
                page_num += 1
        if current:
            pages.append({"page": page_num, "text": "\n".join(current)})
        return pages
    except Exception as e:
        logger.error(f"DOCX parse error for {filename}: {e}")
        return []


def _pdfplumber_page(file_bytes: bytes, page_num: int) -> str:
    """Fallback single-page extraction using pdfplumber."""
    import io
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if page_num < len(pdf.pages):
                return (pdf.pages[page_num].extract_text() or "").strip()
    except Exception as e:
        logger.warning(f"pdfplumber fallback failed page {page_num}: {e}")
    return ""


def _gemini_ocr_pages(file_bytes: bytes, max_pages: int = 8) -> List[Dict]:
    """OCR fallback for scanned PDFs using Gemini multimodal."""
    import io, time
    import fitz
    import google.generativeai as genai
    from config import settings

    genai.configure(api_key=settings.GEMINI_API_KEY)
    model = genai.GenerativeModel(settings.GEMINI_MODEL)

    pages = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total = min(len(doc), max_pages)
        logger.info(f"[OCR] Processing {total} pages via Gemini multimodal...")

        for i in range(total):
            if i > 0:
                time.sleep(3.5)  # avoid rate limit
            page = doc[i]
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")

            import google.generativeai as genai2
            response = model.generate_content([
                "Extract ALL text, tables, and structured content from this document page. "
                "Preserve formatting as markdown. Return only the extracted content.",
                {"mime_type": "image/png", "data": img_bytes},
            ])
            text = response.text.strip() if response.text else ""
            if text:
                pages.append({"page": i + 1, "text": text})
                logger.info(f"[OCR] Page {i+1}/{total} extracted ({len(text)} chars)")
        doc.close()
    except Exception as e:
        logger.error(f"[OCR] Gemini OCR failed: {e}")
    return pages


def parse_pdf(file_bytes: bytes, filename: str) -> List[Dict]:
    """Parse PDF: PyMuPDF → pdfplumber fallback → Gemini OCR fallback."""
    import io
    pages: List[Dict] = []

    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text").strip()
            if len(text) < 50:
                text = _pdfplumber_page(file_bytes, page_num) or text
            if text:
                pages.append({"page": page_num + 1, "text": text})
        doc.close()
    except Exception as e:
        logger.warning(f"PyMuPDF failed for {filename}: {e}")

    if not pages:
        logger.info(f"[Parser] No text extracted from {filename} — trying Gemini OCR")
        pages = _gemini_ocr_pages(file_bytes)

    return pages


def parse_document(file_bytes: bytes, filename: str) -> List[Dict]:
    """Route to the correct parser based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_bytes, filename)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes, filename)
    elif ext in (".txt", ".md"):
        return parse_txt(file_bytes, filename)
    else:
        logger.warning(f"Unsupported file type: {ext}")
        return []
